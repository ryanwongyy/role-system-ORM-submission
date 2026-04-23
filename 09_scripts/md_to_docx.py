#!/usr/bin/env python3
"""Markdown → .docx converter that renders GitHub-flavoured markdown tables
as native Word tables (instead of flattening them to paragraphs like the
older `run_rolemanuscript.write_docx_from_markdown` helper did).

Supports:
- ATX headings (`#`, `##`, ...)
- Paragraphs
- `- ` bullet lists
- `N. ` numbered lists
- GitHub-flavoured `| col | col |` tables with `|---|---|` separator
- Inline bold `**text**` and italic `*text*` (rendered as runs with
  formatting), and backtick `` `code` `` (rendered as run with Consolas)

Usage:
    from md_to_docx import md_to_docx
    md_to_docx(markdown_text, output_path)

Not supported (intentionally kept simple):
- Images, links (rendered as visible text)
- Nested lists, blockquotes, fenced code blocks as multi-line formatted code
- HTML passthrough
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches

DEFAULT_FONT_NAME = "Times New Roman"
DEFAULT_FONT_SIZE = Pt(12)


TABLE_HEADER_RE = re.compile(r"^\s*\|.*\|\s*$")
TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$")
NUMBERED_RE = re.compile(r"^\d+\.\s+")
IMAGE_LINE_RE = re.compile(r"^\s*!\[([^\]]*)\]\(([^)]+)\)\s*$")


def _add_runs_with_inline_formatting(paragraph, text: str) -> None:
    """Parse simple inline markdown and add styled runs to `paragraph`."""
    # Escape order: bold (**...**), italic (*...* or _..._), code (`...`)
    # Use a token scanner; this is intentionally small and fragile — good
    # enough for the manuscript body we control.
    pattern = re.compile(
        r"(\*\*[^*]+\*\*)|(\*[^*]+\*)|(`[^`]+`)|([^*`]+)"
    )
    for match in pattern.finditer(text):
        bold, italic, code, plain = match.groups()
        if bold is not None:
            run = paragraph.add_run(bold[2:-2])
            run.bold = True
        elif italic is not None:
            run = paragraph.add_run(italic[1:-1])
            run.italic = True
        elif code is not None:
            run = paragraph.add_run(code[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10.5)
        elif plain is not None:
            paragraph.add_run(plain)


def _parse_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _add_table(document: Document, header: list[str], body: list[list[str]]) -> None:
    n_cols = len(header)
    table = document.add_table(rows=1 + len(body), cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.autofit = True

    def _style_cell(cell, text: str, bold: bool = False) -> None:
        cell.text = ""
        para = cell.paragraphs[0]
        _add_runs_with_inline_formatting(para, text)
        if bold:
            for run in para.runs:
                run.bold = True

    # Header row
    for j, h in enumerate(header):
        _style_cell(table.rows[0].cells[j], h, bold=True)

    # Body rows
    for i, row in enumerate(body, start=1):
        for j, cell_text in enumerate(row):
            if j < n_cols:
                _style_cell(table.rows[i].cells[j], cell_text)

    # Spacer paragraph after table so the next block renders cleanly
    document.add_paragraph()


def md_to_docx(
    markdown_text: str,
    output_path: Path | str,
    font_name: str = DEFAULT_FONT_NAME,
    font_size: Pt = DEFAULT_FONT_SIZE,
) -> None:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = font_size
    # Set the East Asian font too so the Word picker doesn't fall back.
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)

    lines = markdown_text.splitlines()
    i = 0
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            text = " ".join(part.strip() for part in paragraph_buffer if part.strip())
            if text:
                para = document.add_paragraph()
                _add_runs_with_inline_formatting(para, text)
            paragraph_buffer = []

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        # Markdown table detection
        if (
            TABLE_HEADER_RE.match(line)
            and i + 1 < len(lines)
            and TABLE_SEP_RE.match(lines[i + 1])
        ):
            flush_paragraph()
            header = _parse_table_row(line)
            i += 2  # skip the separator row
            body: list[list[str]] = []
            while i < len(lines) and TABLE_HEADER_RE.match(lines[i]):
                body.append(_parse_table_row(lines[i]))
                i += 1
            _add_table(document, header, body)
            continue

        if not stripped:
            flush_paragraph()
            i += 1
            continue

        if stripped.startswith("#"):
            flush_paragraph()
            level = len(stripped) - len(stripped.lstrip("#"))
            heading_text = stripped[level:].strip()
            # python-docx heading levels: 0 = title, 1..9 = H1..H9
            doc_level = 0 if level == 1 else min(level - 1, 9)
            document.add_heading(heading_text, level=doc_level)
            i += 1
            continue

        if NUMBERED_RE.match(stripped):
            flush_paragraph()
            document.add_paragraph(
                NUMBERED_RE.sub("", stripped), style="List Number"
            )
            i += 1
            continue

        if stripped.startswith("- "):
            flush_paragraph()
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
            i += 1
            continue

        img_match = IMAGE_LINE_RE.match(line)
        if img_match:
            flush_paragraph()
            alt_text, img_path = img_match.group(1), img_match.group(2)
            try:
                # Resolve relative paths against the manuscript location
                from pathlib import Path as _P
                ip = _P(img_path)
                if not ip.is_absolute():
                    ip = (output_path.parent / img_path).resolve()
                if ip.exists():
                    document.add_picture(str(ip), width=Inches(6.0))
                else:
                    document.add_paragraph(f"[Image not found: {img_path}]")
            except Exception as e:
                document.add_paragraph(f"[Image render error: {img_path} — {e}]")
            i += 1
            continue

        paragraph_buffer.append(stripped)
        i += 1

    flush_paragraph()
    document.save(str(output_path))


if __name__ == "__main__":
    # Smoke-test on the manuscript markdown
    ms = Path("/Users/ryanwong/Human roles/p1_role_systems_db/11_manuscript/10_drafts/10_manuscript.md")
    out = ms.with_suffix(".docx")
    md_to_docx(ms.read_text(), out)
    print(f"Wrote {out} ({out.stat().st_size:,} bytes)")
